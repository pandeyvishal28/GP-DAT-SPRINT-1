"""
utils/writing_guide_parser.py
Extracts text content from writing guide files (.txt, .md, .docx, .pdf).

Output is always a UTF-8 string (Markdown-formatted where possible)
ready for LLM prompt injection.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)

# Supported extensions and their handler names (for error messages)
_SUPPORTED_EXTENSIONS: set[str] = {".txt", ".md", ".docx", ".doc", ".pdf"}

# ── PDF font-size thresholds ────────────────────────────────────────────────
_SIZE_H1 = 13.5       # main section headings
_SIZE_H2 = 11.5       # sub-headings
_SIZE_BODY = 10.0     # body text

# Y-coordinate thresholds (PDF points) for header/footer zones
_HEADER_MAX_Y = 100
_FOOTER_MIN_Y = 770


class ParsingError(Exception):
    """Raised when a file cannot be parsed."""


# ── Public API ──────────────────────────────────────────────────────────────


def extract_text(file_path: Path) -> str:
    """
    Extract text content from a writing guide file.

    Parameters
    ----------
    file_path : Path
        Absolute or relative path to the source file.

    Returns
    -------
    str
        Extracted content as a UTF-8 string (Markdown where applicable).

    Raises
    ------
    ParsingError
        If the file type is unsupported or extraction fails.
    FileNotFoundError
        If the file does not exist on disk.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_path.suffix.lower()

    if ext not in _SUPPORTED_EXTENSIONS:
        raise ParsingError(
            f"Unsupported file type '{ext}'. "
            f"Allowed: {', '.join(sorted(_SUPPORTED_EXTENSIONS))}"
        )

    handler = _HANDLERS.get(ext)
    if handler is None:
        raise ParsingError(f"No handler registered for '{ext}'")

    logger.info("Parsing writing guide file: %s (type=%s)", file_path.name, ext)
    return handler(file_path)


def extract_title(md_content: str) -> str | None:
    """
    Extract the document title from Markdown content.
    """
    lines = md_content.splitlines()
    # Labels commonly found on cover pages
    _SKIP_LABELS = {
        "corporate",
        "procedure",
        "sop",
        "work instruction",
        "standard operating procedure",
        "writing guide",
        "general information",
        "type",
        "type:",
        "table of content",
        "table of contents",
        "contents",
        "revision history",
        "document history",
        "change history",
    }

    # Priority: scan for a cover-page metadata table row
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        parts = [p.strip() for p in stripped.strip("|").split("|")]
        if len(parts) >= 2:
            key = parts[0].lower().rstrip(":")
            value = parts[1].strip()
            if key == "title" and value and len(value) > 3:
                if not re.match(r"^[-:]+$", value):
                    return value

    prev_is_title_label = False
    first_real_heading: str | None = None

    for line in lines:
        stripped = line.strip()

        if prev_is_title_label and stripped.startswith("#"):
            return stripped.lstrip("#").strip()

        prev_is_title_label = bool(
            re.search(r"\btitle\b", stripped, re.IGNORECASE) and len(stripped) < 30
        )

        if first_real_heading is None and stripped.startswith("#"):
            text = stripped.lstrip("#").strip()
            if text.lower() not in _SKIP_LABELS and len(text) > 3:
                first_real_heading = text

    return first_real_heading


def extract_description(md_content: str) -> str | None:
    """
    Extract the body text under a 'Purpose' heading.
    """
    lines = md_content.splitlines()
    capture = False
    purpose_level = 0
    parts: list[str] = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped.lstrip("#").strip()

            if re.search(r"\bpurpose\b", heading_text, re.IGNORECASE):
                capture = True
                purpose_level = level
                continue

            if capture and level <= purpose_level:
                break
            continue

        if capture and stripped:
            parts.append(stripped)

    return " ".join(parts).strip() or None


# ── Private handlers ────────────────────────────────────────────────────────


def _parse_text(file_path: Path) -> str:
    """Handle .txt and .md files — plain UTF-8 read."""
    return file_path.read_text(encoding="utf-8")


def _parse_docx(file_path: Path) -> str:
    """
    Handle .docx files using python-docx.

    Extracts paragraphs preserving basic structure:
    - Heading styles → Markdown headings
    - Normal paragraphs → plain text lines
    - Tables → Markdown tables
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ParsingError(
            "python-docx is required for .docx parsing. "
            "Install it with: pip install python-docx"
        ) from exc

    doc = Document(str(file_path))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style else "Normal"
        style_name = style.lower()

        # Map Word heading styles to Markdown headings
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif "heading" in style_name:
            lines.append(f"#### {text}")
        else:
            lines.append(text)

    # Also extract tables
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            md_table = _rows_to_md_table(rows)
            lines.append(md_table)

    return "\n\n".join(lines)


def _parse_pdf(file_path: Path) -> str:
    """
    Handle .pdf files using PyMuPDF and pdfplumber.

    Produces high-quality Markdown output preserving headings,
    tables, and basic formatting based on font size analysis.
    """
    try:
        import fitz  # type: ignore
        import pdfplumber
    except ImportError as exc:
        raise ParsingError(
            "PyMuPDF (fitz) and pdfplumber are required for .pdf parsing. "
            "Install with: pip install PyMuPDF pdfplumber"
        ) from exc

    try:
        # Step 1: extract tables per page via pdfplumber
        tables_by_page: dict[int, list[str]] = {}
        with pdfplumber.open(str(file_path)) as plumber:
            for pg_idx, page in enumerate(plumber.pages):
                md_tables: list[str] = []
                for table in page.extract_tables():
                    md_tables.append(_table_to_markdown(table))
                if md_tables:
                    tables_by_page[pg_idx] = md_tables

        # Step 2: extract text structure via PyMuPDF
        md_parts: list[str] = []
        doc = fitz.open(str(file_path))

        for pg_idx in range(len(doc)):
            page = doc[pg_idx]
            page_h = page.rect.height  # type: ignore
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]  # type: ignore

            page_lines: list[str] = []

            for block in blocks:
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    spans = line["spans"]
                    if not spans:
                        continue

                    # Use first span for font metrics
                    first = spans[0]
                    y_top = first["origin"][1]

                    # Skip header / footer zones
                    if y_top < _HEADER_MAX_Y or y_top > min(_FOOTER_MIN_Y, page_h - 70):
                        continue

                    line_text = "".join(s["text"] for s in spans).strip()
                    if not line_text:
                        continue

                    # Skip single-character watermark artefacts
                    if len(line_text) == 1 and line_text.isalpha():
                        continue

                    size = first["size"]
                    flags = first["flags"]
                    is_bold = bool(flags & 2 ** 4)  # bit 4 = bold

                    md_line = _classify_line(line_text, size, is_bold)
                    page_lines.append(md_line)

            # Append text lines for this page
            if page_lines:
                md_parts.append("\n\n".join(page_lines))

            # Append any tables extracted from this page
            if pg_idx in tables_by_page:
                for tbl_md in tables_by_page[pg_idx]:
                    md_parts.append(tbl_md)

        doc.close()
        return "\n\n".join(md_parts)
    except Exception as exc:
        raise ParsingError(f"Failed to parse PDF '{file_path.name}': {exc}") from exc


# ── Module-level helpers ─────────────────────────────────────────────────


def _classify_line(text: str, size: float, is_bold: bool) -> str:
    """Map a PDF text line to a Markdown line based on font metrics."""
    if is_bold and size >= _SIZE_H1:
        return f"# {text}"
    if is_bold and size >= _SIZE_H2:
        # Determine sub-level from section number pattern
        if re.match(r"^\d+\.\d+\.\d+", text):
            return f"### {text}"
        return f"## {text}"
    return text


def _table_to_markdown(table: list[list[str | None]]) -> str:
    """Convert a pdfplumber raw table (list of rows) to Markdown."""
    if not table:
        return ""
    rows = [[str(cell).strip() if cell else "" for cell in row] for row in table]
    return _rows_to_md_table(rows)


def _rows_to_md_table(rows: list[list[str]]) -> str:
    """Convert a list of string rows into a Markdown table."""
    if not rows:
        return ""
    # Normalise column count
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
    body_lines = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, sep, *body_lines])


# ── Handler registry ────────────────────────────────────────────────────────


_HANDLERS: dict[str, Callable[[Path], str]] = {
    ".txt": _parse_text,
    ".md": _parse_text,
    ".docx": _parse_docx,
    ".doc": _parse_docx,
    ".pdf": _parse_pdf,
}
