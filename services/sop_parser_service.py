"""
services/sop_parser_service.py
Parses SOP documents (PDF, DOCX, TXT) into structured Markdown.

PDF parsing uses a combined approach:
  - PyMuPDF (fitz)  → font-aware heading detection (size, bold)
  - pdfplumber      → table extraction

Heading size thresholds calibrated against BI SOP PDFs
(see data/sop_parser.py for reference values).
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

# ── PDF font-size thresholds (from data/sop_parser.py) ──────────────────
_SIZE_H1 = 13.5       # main section headings: "1 PURPOSE", "2 SCOPE"
_SIZE_H2 = 11.5       # sub-headings: "6.1.1 Risk Based Approach …"
_SIZE_BODY = 10.0      # body text

# Y-coordinate thresholds (PDF points) for header/footer zones
_HEADER_MAX_Y = 100
_FOOTER_MIN_Y = 770


# ── Public API ───────────────────────────────────────────────────────────

class SopParserService:
    """Converts SOP files to Markdown and extracts metadata from the result."""

    # ── Dispatcher ───────────────────────────────────────────────────

    def parse_to_markdown(self, file_path: Path | str, file_type: str) -> str | None:
        """
        Parse a file into Markdown.

        Args:
            file_path: Path to the saved file on disk.
            file_type: One of 'pdf', 'word', 'txt'.

        Returns:
            Markdown string, or None if parsing fails.
        """
        file_path = Path(file_path)
        try:
            if file_type == "pdf":
                return self._parse_pdf_to_markdown(file_path)
            if file_type in ("word", "docx", "doc"):
                return self._parse_docx_to_markdown(file_path)
            if file_type == "txt":
                return self._parse_txt_to_markdown(file_path)
            logger.warning("Unknown file_type '%s'; skipping parse", file_type)
            return None
        except Exception as exc:
            logger.warning("Failed to parse %s to Markdown: %s", file_path.name, exc)
            return None

    # ── Metadata extraction from Markdown ────────────────────────────

    @staticmethod
    def extract_title(md_content: str) -> str | None:
        """
        Extract the document title from Markdown content.

        Strategy:
          1. Look for a ``# `` heading immediately after a line containing
             'title' (cover-page metadata pattern in BI SOPs).
          2. Fallback: return the first ``# `` heading that is not a
             short metadata label (e.g. 'CORPORATE', 'PROCEDURE').
        """
        lines = md_content.splitlines()
        # Labels commonly found on BI SOP cover pages — not actual titles
        _SKIP_LABELS = {
            "corporate", "procedure", "sop", "work instruction",
            "standard operating procedure",
            "general information", "type", "type:",
            "table of content", "table of contents", "contents",
            "revision history", "document history", "change history",
        }

        # ── Priority: scan for a cover-page metadata table row ──────────────
        # BI SOP DOCX files store the title in the first table:
        #   | Title:  | Actual SOP Title |
        # The row appears in the Markdown as a pipe-delimited table line.
        for line in lines:
            stripped = line.strip()
            if not stripped.startswith("|"):
                continue
            parts = [p.strip() for p in stripped.strip("|").split("|")]
            if len(parts) >= 2:
                key = parts[0].lower().rstrip(":")
                value = parts[1].strip()
                if key == "title" and value and len(value) > 3:
                    # Skip separator rows (e.g. | --- | --- |)
                    if not re.match(r"^[-:]+$", value):
                        return value

        prev_is_title_label = False
        first_real_heading: str | None = None

        for line in lines:
            stripped = line.strip()

            # Check if previous line was a "Title:" label
            if prev_is_title_label and stripped.startswith("# ") and not stripped.startswith("## "):
                return stripped.lstrip("# ").strip()

            # Detect "Title:" label lines (heading or plain)
            prev_is_title_label = bool(
                re.search(r"\btitle\b", stripped, re.IGNORECASE)
                and len(stripped) < 30
            )

            # Collect first real heading as fallback
            if (
                first_real_heading is None
                and stripped.startswith("# ")
                and not stripped.startswith("## ")
            ):
                text = stripped.lstrip("# ").strip()
                if text.lower() not in _SKIP_LABELS and len(text) > 3:
                    first_real_heading = text

        return first_real_heading

    @staticmethod
    def extract_description(md_content: str) -> str | None:
        """
        Extract the body text under a 'Purpose' heading.

        Scans for any heading whose text contains the word *purpose*
        (case-insensitive) and collects all non-heading lines until
        the next heading of equal or higher level.
        """
        lines = md_content.splitlines()
        capture = False
        purpose_level = 0
        parts: list[str] = []

        for line in lines:
            stripped = line.strip()

            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                heading_text = stripped.lstrip("#").strip()

                if re.search(r"\bpurpose\b", heading_text, re.IGNORECASE):
                    capture = True
                    purpose_level = level
                    continue

                if capture and level <= purpose_level:
                    break
                continue

            if capture and stripped:
                parts.append(stripped)

        return " ".join(parts).strip() or None

    # ── PDF → Markdown (PyMuPDF + pdfplumber) ────────────────────────

    def _parse_pdf_to_markdown(self, file_path: Path) -> str:
        import fitz          # PyMuPDF
        import pdfplumber

        # ---- Step 1: extract tables per page via pdfplumber ----
        tables_by_page: dict[int, list[str]] = {}
        with pdfplumber.open(str(file_path)) as plumber:
            for pg_idx, page in enumerate(plumber.pages):
                md_tables: list[str] = []
                for table in page.extract_tables():
                    md_tables.append(self._table_to_markdown(table))
                if md_tables:
                    tables_by_page[pg_idx] = md_tables

        # ---- Step 2: extract text structure via PyMuPDF ----
        md_parts: list[str] = []
        doc = fitz.open(str(file_path))

        for pg_idx in range(len(doc)):
            page = doc[pg_idx]
            page_h = page.rect.height
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

            page_lines: list[str] = []

            for block in blocks:
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    spans = line["spans"]
                    if not spans:
                        continue

                    # Use first span for font metrics
                    first = spans[0]
                    y_top = first["origin"][1]

                    # Skip header / footer zones
                    if y_top < _HEADER_MAX_Y or y_top > min(_FOOTER_MIN_Y, page_h - 70):
                        continue

                    line_text = "".join(s["text"] for s in spans).strip()
                    if not line_text:
                        continue

                    # Skip single-character watermark artefacts
                    if len(line_text) == 1 and line_text.isalpha():
                        continue

                    size = first["size"]
                    flags = first["flags"]
                    is_bold = bool(flags & 2 ** 4)  # bit 4 = bold

                    md_line = self._classify_line(line_text, size, is_bold)
                    page_lines.append(md_line)

            # Append text lines for this page
            if page_lines:
                md_parts.append("\n\n".join(page_lines))

            # Append any tables extracted from this page
            if pg_idx in tables_by_page:
                for tbl_md in tables_by_page[pg_idx]:
                    md_parts.append(tbl_md)

        doc.close()
        return "\n\n".join(md_parts)

    # ── DOCX → Markdown ──────────────────────────────────────────────

    @staticmethod
    def _parse_docx_to_markdown(file_path: Path) -> str:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name:
                lines.append(f"# {text}")
            elif "heading 2" in style_name:
                lines.append(f"## {text}")
            elif "heading 3" in style_name:
                lines.append(f"### {text}")
            elif "heading" in style_name:
                lines.append(f"#### {text}")
            else:
                lines.append(text)

        # Also extract tables
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                md_table = _rows_to_md_table(rows)
                lines.append(md_table)

        return "\n\n".join(lines)

    # ── TXT → Markdown ───────────────────────────────────────────────

    @staticmethod
    def _parse_txt_to_markdown(file_path: Path) -> str:
        return file_path.read_text(encoding="utf-8", errors="replace")

    # ── Helpers ──────────────────────────────────────────────────────

    @staticmethod
    def _classify_line(text: str, size: float, is_bold: bool) -> str:
        """Map a PDF text line to a Markdown line based on font metrics."""
        if is_bold and size >= _SIZE_H1:
            return f"# {text}"
        if is_bold and size >= _SIZE_H2:
            # Determine sub-level from section number pattern
            if re.match(r"^\d+\.\d+\.\d+", text):
                return f"### {text}"
            return f"## {text}"
        return text

    @staticmethod
    def _table_to_markdown(table: list[list[str | None]]) -> str:
        """Convert a pdfplumber raw table (list of rows) to Markdown."""
        if not table:
            return ""
        rows = [[str(cell).strip() if cell else "" for cell in row] for row in table]
        return _rows_to_md_table(rows)


# ── Module-level helpers ─────────────────────────────────────────────────

def _rows_to_md_table(rows: list[list[str]]) -> str:
    """Convert a list of string rows into a Markdown table."""
    if not rows:
        return ""
    # Normalise column count
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
    body_lines = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, sep, *body_lines])
