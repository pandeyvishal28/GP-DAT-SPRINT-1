"""
services/pdf_converter.py
PDF → DOCX conversion via pdf2docx, with light post-processing.
"""

from __future__ import annotations

import copy
import logging
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement

# pdf2docx 0.5.8 expects .part on XML elements (removed in python-docx >= 1.0).
# Restore the old behaviour: store _part on the root, walk up via getparent().
if not hasattr(BaseOxmlElement, "part"):

    def _element_part(self):
        e = self
        while e is not None:
            if hasattr(e, "_part"):
                return e._part
            e = e.getparent()
        raise AttributeError(f"{type(self).__name__} element is not in a Document")

    BaseOxmlElement.part = property(_element_part)

    import docx.document as _docmod

    _orig_init = _docmod.Document.__init__

    def _patched_init(self, *args, **kwargs):
        _orig_init(self, *args, **kwargs)
        self.element._part = self.part

    _docmod.Document.__init__ = _patched_init

from pdf2docx import Converter

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path("data/outputs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ─── Public API ─────────────────────────────────────────────────────


def convert_pdf_to_docx(pdf_path: Path, output_dir: Path | None = None) -> Path:
    out = (output_dir or OUTPUT_DIR) / f"Converted_{pdf_path.stem}_{_ts()}.docx"
    logger.info("PDF → DOCX: %s", pdf_path.name)

    cv = Converter(str(pdf_path))
    try:
        cv.convert(str(out))
    finally:
        cv.close()

    _postprocess(out)
    logger.info("Conversion done: %s (%d bytes)", out.name, out.stat().st_size)
    return out


def prepare_working_docx(
    file_path: Path, file_ext: str, output_dir: Path | None = None,
) -> Path:
    dest = output_dir or OUTPUT_DIR
    dest.mkdir(parents=True, exist_ok=True)

    if file_ext == ".pdf":
        return convert_pdf_to_docx(file_path, dest)

    if file_ext in (".docx", ".doc"):
        out = dest / f"WorkingCopy_{file_path.stem}_{_ts()}.docx"
        shutil.copy2(file_path, out)
        return out

    out = dest / f"WorkingCopy_{file_path.stem}_{_ts()}.docx"
    doc = Document()
    for line in file_path.read_text(encoding="utf-8", errors="replace").splitlines():
        doc.add_paragraph(line)
    doc.save(str(out))
    return out


# ─── Post-processing ───────────────────────────────────────────────


def _postprocess(docx_path: Path) -> None:
    try:
        doc = Document(str(docx_path))
        _populate_header_footer(doc)
        _fix_margins(doc)
        _split_toc_paragraph(doc)
        _fix_heading_spacing(doc)
        _strip_images_after_page1(doc)
        _trim_trailing_blanks(doc)
        doc.save(str(docx_path))
    except Exception as e:
        logger.warning("Post-processing skipped: %s", e)


# ─── Compact boilerplate spacing ──────────────────────────────────

_BP_FOOTER = re.compile(
    r"(Property of|Proprietary confidential.*?)Boehringer Ingelheim",
    re.IGNORECASE,
)
_BP_RUNNING_HDR = re.compile(
    r"(Nummer|Number):\s*\w[\w-]+.*Version:", re.IGNORECASE,
)
_BP_PAGE_NUM = re.compile(
    r"(Seite|Page)\s+\d+\s+(von|of)\s+\d+", re.IGNORECASE,
)
_BP_RETRIEVED = re.compile(r"Retrieved by\b", re.IGNORECASE)
_BP_VERIFY_VER = re.compile(r"Verify the current version", re.IGNORECASE)


# ─── Populate proper DOCX header / footer sections ────────────────


def _populate_header_footer(doc: Document) -> None:
    """Copy boilerplate into proper DOCX header/footer sections,
    then remove the duplicates from the body.
    """
    body = doc.element.body
    _T = qn("w:t")
    _P = qn("w:p")
    _TBL = qn("w:tbl")
    _TC = qn("w:tc")

    def _txt(el):
        return "".join(t.text or "" for t in el.findall(".//" + _T)).strip()

    def _is_hdr_tbl(el):
        if el.tag != _TBL:
            return False
        cells = el.findall(".//" + _TC)
        full = " ".join(_txt(c) for c in cells)
        return "Boehringer Ingelheim" in full and (
            "Document Name" in full or "Dokumentname" in full
            or "Document Title" in full or "Document ID" in full
        )

    def _is_ftr_tbl(el):
        """Detect footer tables: contain footer/page-number text but NOT header-table markers."""
        if el.tag != _TBL:
            return False
        if _is_hdr_tbl(el):
            return False
        cells = el.findall(".//" + _TC)
        full = " ".join(_txt(c) for c in cells)
        return bool(
            _BP_FOOTER.search(full) or _BP_PAGE_NUM.search(full)
            or _BP_RETRIEVED.search(full) or _BP_VERIFY_VER.search(full)
        )

    def _is_ftr_para(t):
        """Return True if text matches any footer boilerplate pattern."""
        return bool(
            _BP_FOOTER.search(t) or _BP_PAGE_NUM.search(t)
            or _BP_RETRIEVED.search(t) or _BP_VERIFY_VER.search(t)
        )

    # Collect one footer paragraph per pattern type (avoids duplicates from
    # combined paragraphs vs. individual ones across pages).
    _FTR_PATTERNS = [
        ("proprietary", _BP_FOOTER),
        ("retrieved", _BP_RETRIEVED),
        ("verify", _BP_VERIFY_VER),
    ]

    hdr_txt_el = None
    hdr_tbl_el = None
    ftr_txt_els = []
    ftr_tbl_el = None
    _ftr_found = set()        # pattern names already collected

    def _collect_ftr(el):
        """Collect a footer paragraph if it has any new pattern type."""
        t = _txt(el)
        types = [name for name, pat in _FTR_PATTERNS if pat.search(t)]
        new_types = [name for name in types if name not in _ftr_found]
        if not new_types:
            return
        ftr_txt_els.append(el)
        _ftr_found.update(types)   # mark ALL matched types as found

    for el in list(body):
        if el.tag == _P:
            t = _txt(el)
            if _BP_RUNNING_HDR.search(t) and hdr_txt_el is None:
                hdr_txt_el = el
            elif _is_ftr_para(t):
                _collect_ftr(el)
        elif _is_hdr_tbl(el) and hdr_tbl_el is None:
            hdr_tbl_el = el
        elif _is_ftr_tbl(el) and ftr_tbl_el is None:
            ftr_tbl_el = el

    # If no footer paragraphs found at top level, search inside tables
    if not ftr_txt_els:
        for p_elem in body.iter(_P):
            t = _txt(p_elem)
            if _is_ftr_para(t):
                _collect_ftr(p_elem)

    # Populate header on section 0 (sections 1+ inherit via linked_to_previous)
    if hdr_txt_el is not None or hdr_tbl_el is not None:
        sec0 = doc.sections[0]
        sec0.header.is_linked_to_previous = False
        h = sec0.header._element
        for ch in list(h):
            h.remove(ch)
        if hdr_txt_el is not None:
            cp = copy.deepcopy(hdr_txt_el)
            _zero_spacing(cp)
            h.append(cp)
        if hdr_tbl_el is not None:
            h.append(copy.deepcopy(hdr_tbl_el))
        for sec in doc.sections[1:]:
            sec.header.is_linked_to_previous = True

    # Populate footer on section 0 (all others inherit)
    if ftr_txt_els or ftr_tbl_el is not None:
        sec0 = doc.sections[0]
        sec0.footer.is_linked_to_previous = False
        f = sec0.footer._element
        for ch in list(f):
            f.remove(ch)
        for ftr_el in ftr_txt_els:
            cp = copy.deepcopy(ftr_el)
            _zero_spacing(cp)
            f.append(cp)
        if ftr_tbl_el is not None:
            f.append(copy.deepcopy(ftr_tbl_el))
        for sec in doc.sections[1:]:
            sec.footer.is_linked_to_previous = True

    # Remove ALL body copies of the boilerplate (now in header/footer)
    removed = 0
    for el in list(body):
        if el.tag == _P:
            t = _txt(el)
            if (
                _BP_RUNNING_HDR.search(t) or _BP_FOOTER.search(t)
                or _BP_PAGE_NUM.search(t) or _BP_RETRIEVED.search(t)
                or _BP_VERIFY_VER.search(t)
            ):
                body.remove(el)
                removed += 1
        elif _is_hdr_tbl(el):
            body.remove(el)
            removed += 1
        elif _is_ftr_tbl(el):
            body.remove(el)
            removed += 1

    if removed:
        logger.info("Removed %d body boilerplate duplicates", removed)


def _zero_spacing(p_elem) -> None:
    """Set before/after to 0 and line spacing to single on a paragraph."""
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        return
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "auto")


# ─── Margin fix ─────────────────────────────────────────────────────

# Original PDF: top ~0.22cm, bottom ~0.45cm
# We use slightly larger values to leave room for the DOCX header/footer
_HEADER_DISTANCE = 79_200    # ~0.22cm — matches PDF top text start
_FOOTER_DISTANCE = 162_000   # ~0.45cm — matches PDF bottom text end
_TOP_MARGIN = 360_000        # ~1cm — body starts below header content
_BOTTOM_MARGIN = 180_000     # ~0.5cm — body ends above footer content


def _fix_margins(doc: Document) -> None:
    """Set margins and header/footer distances to match the original PDF."""
    for sec in doc.sections:
        sec.top_margin = _TOP_MARGIN
        sec.bottom_margin = _BOTTOM_MARGIN
        sec.header_distance = _HEADER_DISTANCE
        sec.footer_distance = _FOOTER_DISTANCE


# ─── TOC paragraph splitting ──────────────────────────────────────

_TOC_ENTRY_RE = re.compile(
    r"(\d+(?:\.\d+)*)\s+"   # section number
    r"(.+?)"                 # title (non-greedy)
    r"\s*\.{4,}\s*"          # dot leader
    r"(\d{1,3})"             # page number
)


def _split_toc_paragraph(doc: Document) -> None:
    """Detect a mega-TOC paragraph (all entries in one) and split it."""
    body = doc.element.body
    _T = qn("w:t")

    for p_elem in list(body):
        if p_elem.tag != qn("w:p"):
            continue
        full_text = "".join(t.text or "" for t in p_elem.findall(".//" + _T))
        entries = _TOC_ENTRY_RE.findall(full_text)
        if len(entries) < 3:
            continue

        # Copy source formatting
        pPr_src = p_elem.find(qn("w:pPr"))
        r_src = p_elem.find(qn("w:r"))
        rPr_src = r_src.find(qn("w:rPr")) if r_src is not None else None

        for sec_num, title, page_num in entries:
            new_p = OxmlElement("w:p")
            if pPr_src is not None:
                new_pPr = copy.deepcopy(pPr_src)
                sp = new_pPr.find(qn("w:sectPr"))
                if sp is not None:
                    new_pPr.remove(sp)
                new_p.append(new_pPr)

            new_r = OxmlElement("w:r")
            if rPr_src is not None:
                new_r.append(copy.deepcopy(rPr_src))

            t_el = OxmlElement("w:t")
            t_el.text = f"{sec_num} {title.strip()}"
            t_el.set(qn("xml:space"), "preserve")
            new_r.append(t_el)

            new_r.append(OxmlElement("w:tab"))

            pn = OxmlElement("w:t")
            pn.text = page_num
            pn.set(qn("xml:space"), "preserve")
            new_r.append(pn)

            new_p.append(new_r)
            p_elem.addprevious(new_p)

        body.remove(p_elem)
        logger.info("Split TOC into %d entries", len(entries))
        break


# ─── Heading spacing fix ───────────────────────────────────────────


def _fix_heading_spacing(doc: Document) -> None:
    """Turn '1.Introduction' into '1. Introduction'."""
    for para in doc.paragraphs:
        if not para.style.name.startswith("Heading"):
            continue
        runs = para.runs
        if len(runs) < 2:
            continue
        first = runs[0].text.strip()
        if first and first[-1] == "." and first[:-1].replace(".", "").isdigit():
            if runs[1].text and not runs[1].text.startswith(" "):
                runs[1].text = " " + runs[1].text


def _strip_images_after_page1(doc: Document) -> None:
    """Keep first-page images only; remove all others."""
    past_page1 = False
    removed = 0
    _BR = qn("w:br")
    _SECT = qn("w:sectPr")
    _DRAW = qn("w:drawing")
    _PICT = qn("w:pict")

    for elem in list(doc.element.body.iter()):
        if not past_page1:
            if elem.tag == _BR and elem.get(qn("w:type")) == "page":
                past_page1 = True
            elif elem.tag == _SECT:
                past_page1 = True
            continue

        if elem.tag in (_DRAW, _PICT):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                removed += 1

    if removed:
        logger.info("Stripped %d images after page 1", removed)


def _trim_trailing_blanks(doc: Document) -> None:
    """Remove empty trailing paragraphs that create blank last pages."""
    body = doc.element.body
    _P = qn("w:p")
    _T = qn("w:t")
    _SECT = qn("w:sectPr")

    for elem in reversed(list(body)):
        if elem.tag == _SECT:
            continue
        if elem.tag == _P:
            text = "".join(t.text or "" for t in elem.iter(_T))
            if not text.strip():
                body.remove(elem)
                continue
        break


def _ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")
