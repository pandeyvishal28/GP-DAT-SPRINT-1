"""
GP Docs Template Parser
Converts a structured .docx template into a clean Markdown file for LLM consumption.

Usage:
    python parse_template.py <input.docx> <output.md>

Output element types:
    CHAPTER_TITLE   - Heading 1 style  → becomes # heading in MD
    CONTENT_FIXED   - Black text        → verbatim in output, marked clearly
    PLACEHOLDER     - ${vault:…} fields → becomes {{field_name}} in MD
    INSTRUCTION     - Blue text /       → becomes [!INSTRUCTION] callout
                      annotation style
    TABLE           - parsed cell by    → Markdown table with per-cell types
                      cell
"""

import logging
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# ── Constants ────────────────────────────────────────────────────────────────

BLUE_COLOR = "0075FF"  # The exact blue used for all instructions
PLACEHOLDER_RE = re.compile(r"\$\{vault:([^}]+)\}")  # matches ${vault:field__v}
DIVIDER_RE = re.compile(r"^\*+$")  # matches *** separator lines

# Styles that are always instructions regardless of color
INSTRUCTION_STYLES = {
    "annotation text",
    "Footer",
}

# Style name prefixes that are always instructions
INSTRUCTION_PREFIXES = ("Instructions",)

# Styles that are always chapter titles
HEADING_STYLES = {
    "Heading 1",
}


# ── Color helpers ─────────────────────────────────────────────────────────────


def get_run_colors(paragraph) -> set:
    """
    Extract all explicit font colors set on runs within a paragraph.
    Returns a set of hex color strings e.g. {'0075FF', 'auto'}.
    'auto' means the run inherits the theme/default color (black).
    A run with no <w:color> element also inherits (treated as black).
    """
    colors = set()
    for run in paragraph.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            color_el = rpr.find(qn("w:color"))
            if color_el is not None:
                val = color_el.get(qn("w:val"))
                if val:
                    colors.add(val.upper())
    return colors


def is_instruction_color(colors: set) -> bool:
    """
    Returns True if the paragraph contains any blue instruction runs.
    A paragraph is instruction if ANY run is blue — even mixed paragraphs.
    """
    return BLUE_COLOR in colors


def get_cell_colors(cell) -> set:
    """
    Aggregate run colors across all paragraphs in a table cell.
    """
    colors = set()
    for para in cell.paragraphs:
        colors |= get_run_colors(para)
    return colors


# ── Placeholder helpers ───────────────────────────────────────────────────────


def extract_placeholders(text: str) -> list:
    """
    Find all ${vault:field_name} patterns in text.
    Returns list of clean field names e.g. ['impacted_divisions__c', 'type__v']
    """
    return PLACEHOLDER_RE.findall(text)


def normalize_placeholder(field_name: str) -> str:
    """
    Convert vault field names to clean {{placeholder}} tokens for Markdown.
    e.g. 'impacted_divisions__c' → '{{impacted_divisions}}'
         'major_version_number__v' → '{{major_version_number}}'
    """
    # Strip Veeva suffixes __c (custom) and __v (vault standard)
    clean = re.sub(r"__(c|v)$", "", field_name)
    return "{{" + clean + "}}"


def replace_placeholders(text: str) -> str:
    """
    Replace all ${vault:field_name} tokens with clean {{field_name}} tokens.
    """

    def replacer(match):
        return normalize_placeholder(match.group(1))

    return PLACEHOLDER_RE.sub(replacer, text)


# ── Paragraph classifier ──────────────────────────────────────────────────────


def classify_paragraph(para) -> dict | None:
    """
    Classify a single paragraph into one of five element types.
    Returns a dict with 'type' and 'text', or None if the paragraph
    should be completely discarded (empty, divider, etc).

    Classification order matters — each rule short-circuits:
      1. Discard  — empty text or visual dividers
      2. INSTRUCTION — style name matches known instruction styles/prefixes
      3. CHAPTER_TITLE — Heading 1
      4. INSTRUCTION — color check (catches Normal paragraphs styled blue)
      5. PLACEHOLDER — contains ${vault:…} pattern
      6. CONTENT_FIXED — everything else with actual text
    """
    text = para.text.strip()

    # ── Rule 1: Discard empty and divider paragraphs ──────────────────────────
    if not text:
        return None
    if DIVIDER_RE.match(text):
        return None

    style = para.style.name

    # ── Rule 2: Style-based instruction detection ─────────────────────────────
    if style in INSTRUCTION_STYLES:
        return {"type": "INSTRUCTION", "text": text}
    if style.startswith(INSTRUCTION_PREFIXES):
        return {"type": "INSTRUCTION", "text": text}

    # ── Rule 3: Chapter title ─────────────────────────────────────────────────
    if style in HEADING_STYLES:
        return {"type": "CHAPTER_TITLE", "text": text}

    # ── Rule 4: Color-based instruction detection ─────────────────────────────
    # This catches paragraphs in Normal/List/Text styles that are blue.
    # Must come AFTER heading check — headings are never blue but we
    # want to be safe about order.
    colors = get_run_colors(para)
    if is_instruction_color(colors):
        return {"type": "INSTRUCTION", "text": text}

    # ── Rule 5: Placeholder detection ────────────────────────────────────────
    if PLACEHOLDER_RE.search(text):
        clean_text = replace_placeholders(text)
        fields = extract_placeholders(text)
        return {
            "type": "PLACEHOLDER",
            "text": clean_text,
            "fields": [normalize_placeholder(f) for f in fields],
        }

    # ── Rule 6: Fixed content ─────────────────────────────────────────────────
    return {"type": "CONTENT_FIXED", "text": text}


# ── Table classifier ──────────────────────────────────────────────────────────


def classify_cell(cell) -> dict | None:
    """
    Classify a single table cell using the same color/placeholder logic
    as paragraph classification, but operating across all paragraphs in the cell.

    Returns dict with 'type' and 'text', or None if cell is empty.
    """
    # Collect full cell text across all paragraphs
    full_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())

    if not full_text:
        return None

    # Aggregate colors across all runs in all paragraphs of this cell
    colors = get_cell_colors(cell)

    # Blue → instruction
    if is_instruction_color(colors):
        return {"type": "INSTRUCTION", "text": full_text}

    # Placeholder pattern
    if PLACEHOLDER_RE.search(full_text):
        clean = replace_placeholders(full_text)
        fields = extract_placeholders(full_text)
        return {
            "type": "PLACEHOLDER",
            "text": clean,
            "fields": [normalize_placeholder(f) for f in fields],
        }

    # Everything else is fixed content
    return {"type": "CONTENT_FIXED", "text": full_text}


def classify_table(table) -> dict:
    """
    Classify an entire table by processing it cell by cell.

    Determines table-level type:
    - 'metadata'     : first column has fixed labels, second has placeholders
                       (the cover page info table)
    - 'schema'       : header row is fixed, body rows are instructions
                       (template tables the LLM should fill)
    - 'instruction'  : all cells are instructions — entire table is guidance only
    - 'mixed'        : combination (render with per-cell markers)

    Also extracts:
    - headers: list of header cell texts (first row)
    - rows: list of classified rows
    """
    classified_rows = []

    for row in table.rows:
        classified_row = []
        for cell in row.cells:
            classified_row.append(classify_cell(cell))
        classified_rows.append(classified_row)

    # Determine overall table type
    # Check if ALL non-empty cells are instructions
    all_cells = [c for row in classified_rows for c in row if c is not None]
    if not all_cells:
        return {"type": "empty", "rows": classified_rows}

    all_instruction = all(c["type"] == "INSTRUCTION" for c in all_cells)
    if all_instruction:
        return {"type": "instruction", "rows": classified_rows}

    # Check if it looks like a metadata table (col 0 = fixed label, col 1 = placeholder)
    has_placeholders = any(c["type"] == "PLACEHOLDER" for c in all_cells)
    has_fixed = any(c["type"] == "CONTENT_FIXED" for c in all_cells)

    if has_placeholders and has_fixed:
        # Check if it's a clean key-value metadata table
        # (all rows have fixed in col 0, placeholder in col 1)
        is_metadata = all(
            len(row) >= 2
            and row[0] is not None
            and row[0]["type"] == "CONTENT_FIXED"
            and (row[1] is None or row[1]["type"] in ("PLACEHOLDER", "CONTENT_FIXED"))
            for row in classified_rows
            if any(c is not None for c in row)
        )
        if is_metadata:
            return {"type": "metadata", "rows": classified_rows}
        return {"type": "mixed", "rows": classified_rows}

    return {"type": "schema", "rows": classified_rows}


# ── Body element iterator ─────────────────────────────────────────────────────


def iter_body_elements(doc):
    """
    Iterate document body children in document order, yielding
    ('paragraph', para_obj) or ('table', table_obj) tuples.

    This is the ONLY correct way to preserve paragraph/table interleaving.
    doc.paragraphs and doc.tables are separate flat lists that lose order.
    """
    # Map lxml elements back to python-docx objects
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in doc.element.body:
        if child in para_map:
            yield ("paragraph", para_map[child])
        elif child in table_map:
            yield ("table", table_map[child])
        # Other elements (sectPr, etc.) are skipped


# ── Chapter grouper ───────────────────────────────────────────────────────────


def group_into_chapters(doc) -> list:
    """
    Walk the document body in order and group all elements under their
    parent Heading 1 chapter.

    Returns a list of chapter dicts:
    {
        "title": str,
        "chapter_id": str,   # slug for referencing
        "elements": [...]    # classified paragraphs and tables in order
    }

    Elements before the first Heading 1 go into a special "preamble" chapter.
    """
    chapters = []
    current_chapter: dict[str, Any] = {
        "title": "PREAMBLE",
        "chapter_id": "preamble",
        "elements": [],
    }

    for element_type, element in iter_body_elements(doc):
        if element_type == "paragraph":
            classified = classify_paragraph(element)
            if classified is None:
                continue

            if classified["type"] == "CHAPTER_TITLE":
                # Save current chapter if it has content
                if current_chapter["elements"]:
                    chapters.append(current_chapter)
                # Start new chapter
                title = classified["text"]
                chapter_id = re.sub(r"[^a-z0-9]+", "_", title.lower()).strip("_")
                current_chapter = {
                    "title": title,
                    "chapter_id": chapter_id,
                    "elements": [],
                }
            else:
                current_chapter["elements"].append(classified)

        elif element_type == "table":
            classified_table = classify_table(element)
            # Skip fully-instruction tables (they're guidance only)
            # Keep everything else — metadata, schema, mixed
            if classified_table["type"] != "instruction":
                current_chapter["elements"].append(
                    {
                        "type": "TABLE",
                        "table_type": classified_table["type"],
                        "rows": classified_table["rows"],
                    }
                )
            else:
                # Still capture it as instruction context for the LLM
                # Extract text from instruction cells as bullet list
                instruction_texts = []
                for row in classified_table["rows"]:
                    for cell in row:
                        if cell and cell["text"]:
                            instruction_texts.append(cell["text"])
                if instruction_texts:
                    current_chapter["elements"].append(
                        {
                            "type": "INSTRUCTION",
                            "text": "\n".join(instruction_texts),
                        }
                    )

    # Don't forget the last chapter
    if current_chapter["elements"]:
        chapters.append(current_chapter)

    return chapters


# ── Markdown renderer ─────────────────────────────────────────────────────────

# Headers that indicate a table is meta/instructional and should NOT be
# reproduced by the LLM in the output document.
_INSTRUCTIONAL_TABLE_HEADERS = {"infographics"}


def _is_instructional_table(table_element: dict) -> bool:
    """
    Detect tables that are instructional guidance rather than output tables.

    A table is instructional if:
    1. Its first header cell matches a known instructional keyword (e.g. "Infographics"), OR
    2. The majority of its body cells are INSTRUCTION-type cells.
    """
    rows = table_element.get("rows", [])
    if not rows:
        return False

    # Check 1: header keyword
    first_header = rows[0][0]
    if first_header and first_header.get("text", "").strip().lower() in _INSTRUCTIONAL_TABLE_HEADERS:
        return True

    # Check 2: majority of body cells are INSTRUCTION
    body_cells = [cell for row in rows[1:] for cell in row if cell]
    if not body_cells:
        return False
    instruction_count = sum(1 for c in body_cells if c.get("type") == "INSTRUCTION")
    return instruction_count > len(body_cells) * 0.5


def _get_table_label(table_element: dict) -> str:
    """
    Derive a human-readable label for a table from its first-row headers.
    Used to insert <!-- TABLE: Label --> comments so consecutive tables
    don't merge in the LLM's perception.
    """
    rows = table_element.get("rows", [])
    if not rows:
        return "Table"
    headers = [
        c.get("text", "").strip()
        for c in rows[0]
        if c and c.get("text", "").strip()
    ]
    if headers:
        return " / ".join(h[:40] for h in headers[:3])
    return "Table"


def render_table_to_markdown(table_element: dict) -> str:
    """
    Render a classified table to Markdown format.

    - CONTENT_FIXED cells render as plain text
    - PLACEHOLDER cells render as {{field_name}}
    - INSTRUCTION cells render as *italics* with a note
    - Empty cells render as empty string

    Table type affects rendering:
    - metadata: renders as a simple two-column key/value table
    - schema: renders with a note that LLM should populate body rows
    - mixed: renders as-is with inline markers
    """
    rows = table_element["rows"]
    table_type = table_element["table_type"]

    if not rows:
        return ""

    md_rows = []
    for row in rows:
        md_cells = []
        for cell in row:
            if cell is None:
                md_cells.append("")
            elif cell["type"] == "CONTENT_FIXED":
                md_cells.append(cell["text"].replace("\n", " ").replace("|", "\\|"))
            elif cell["type"] == "PLACEHOLDER":
                md_cells.append(cell["text"].replace("\n", " "))
            elif cell["type"] == "INSTRUCTION":
                # Instruction cells: render as italic so LLM knows not to reproduce
                short = cell["text"].split("\n")[0][:80]
                md_cells.append(f"*{short}*")
        md_rows.append(md_cells)

    if not md_rows:
        return ""

    # Build markdown table
    lines = []
    num_cols = max(len(r) for r in md_rows)

    # Pad all rows to same column count
    for row in md_rows:
        while len(row) < num_cols:
            row.append("")

    # Header row (first row)
    header = md_rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * num_cols) + " |")

    # Body rows
    for row in md_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    # Add hint for schema tables so LLM knows to generate rows
    if table_type == "schema":
        lines.append("")
        lines.append(
            "> *Generate additional rows as needed based on the document content.*"
        )

    return "\n".join(lines)


def render_to_markdown(chapters: list) -> str:
    """
    Render the full chapter list to a Markdown string.

    Element type → Markdown mapping:
        INSTRUCTION   → > [!INSTRUCTION] callout block
        CONTENT_FIXED → plain text
        PLACEHOLDER   → plain text with {{tokens}}
        TABLE         → Markdown table (or [!INSTRUCTION] if instructional)
    """
    lines = []

    # File header
    lines.append("# DOCUMENT TEMPLATE: GP Docs")
    lines.append("")
    lines.append("> [!SYSTEM]")
    lines.append("> You are generating a controlled regulatory document.")
    lines.append(
        "> - `[!INSTRUCTION]` callout blocks are directions for you — **never reproduce them in output**."
    )
    lines.append(
        "> - Fixed content text outside of callouts must appear **verbatim** in your output."
    )
    lines.append("> - `{{placeholder}}` tokens are where you generate content.")
    lines.append("> - Maintain formal, regulatory-compliant language throughout.")
    lines.append(
        "> - Each table is labeled with a `<!-- TABLE: ... -->` comment."
        " Keep tables **separate** \u2014 never merge consecutive tables."
    )
    lines.append("")

    for chapter in chapters:
        # Skip preamble if it has no meaningful content
        if chapter["chapter_id"] == "preamble":
            meaningful = [
                e
                for e in chapter["elements"]
                if e["type"] in ("CONTENT_FIXED", "PLACEHOLDER")
            ]
            if not meaningful:
                continue

        # Chapter heading
        if chapter["chapter_id"] != "preamble":
            lines.append(f"# CHAPTER: {chapter['title']}")
            lines.append(f"<!-- chapter_id: {chapter['chapter_id']} -->")
            lines.append("")

        # Collect consecutive instructions to group into one callout block
        # This avoids many tiny separate callout blocks for the same section
        pending_instructions: list[str] = []
        prev_was_table = False  # Track consecutive tables to insert separators

        def flush_instructions():
            if pending_instructions:
                lines.append("> [!INSTRUCTION]")
                for instr in pending_instructions:
                    # Multi-line instructions: each line gets > prefix
                    for sub_line in instr.split("\n"):
                        sub_line = sub_line.strip()
                        if sub_line:
                            lines.append(f"> {sub_line}")
                lines.append("")
                pending_instructions.clear()

        for element in chapter["elements"]:
            etype = element["type"]

            if etype == "INSTRUCTION":
                pending_instructions.append(element["text"])
                prev_was_table = False

            else:
                # Flush any pending instructions before non-instruction content
                flush_instructions()

                if etype == "CONTENT_FIXED":
                    lines.append(f"{element['text']}")
                    lines.append("")
                    prev_was_table = False

                elif etype == "PLACEHOLDER":
                    lines.append(element["text"])
                    lines.append("")
                    prev_was_table = False

                elif etype == "TABLE":
                    # ── Check if this table is instructional / meta ──
                    if _is_instructional_table(element):
                        label = _get_table_label(element)
                        logger.debug("Wrapping instructional table as callout: %s", label)
                        table_md = render_table_to_markdown(element)
                        if table_md:
                            lines.append("> [!INSTRUCTION]")
                            lines.append(
                                "> The following table is for reference only"
                                " \u2014 do NOT include it in the output:"
                            )
                            for tline in table_md.splitlines():
                                lines.append(f"> {tline}")
                            lines.append("")
                        prev_was_table = False  # It's wrapped as instruction now
                        continue

                    # ── Normal table: add label and separator ──
                    label = _get_table_label(element)

                    # Insert a blank line + separator between consecutive tables
                    if prev_was_table:
                        lines.append("")

                    lines.append(f"<!-- TABLE: {label} -->")
                    table_md = render_table_to_markdown(element)
                    if table_md:
                        lines.append(table_md)
                        lines.append("")
                    prev_was_table = True

        # Flush any trailing instructions at end of chapter
        flush_instructions()

    return "\n".join(lines)


# ── Service entry point ───────────────────────────────────────────────────────


def parse_docx_to_blueprint(
    docx_path: str | Path,
    output_dir: str | Path = "data/blueprints",
) -> Path:
    """
    Parse a .docx template and save it as a Markdown blueprint.

    Args:
        docx_path:  Path to the .docx template file.
        output_dir: Directory where the .md blueprint is saved.

    Returns:
        Path to the generated .md blueprint file.

    The output filename is auto-derived from the input:
        'Template Main GP Docs.docx' → 'template_main_gp_docs.md'
    """
    docx_path = Path(docx_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Derive blueprint filename from docx name (matches template_id convention)
    blueprint_name = docx_path.stem.lower().replace(" ", "_") + ".md"
    output_path = output_dir / blueprint_name

    logger.info("Parsing template: %s", docx_path.name)
    doc = Document(str(docx_path))

    chapters = group_into_chapters(doc)
    logger.info(
        "Template grouped into %d chapters from %s",
        len(chapters), docx_path.name,
    )

    markdown = render_to_markdown(chapters)

    output_path.write_text(markdown, encoding="utf-8")
    logger.info(
        "Blueprint saved: %s (%d chars, %d lines)",
        output_path.name, len(markdown), len(markdown.splitlines()),
    )

    return output_path


# ── CLI fallback ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(message)s")
    if len(sys.argv) < 2:
        print("Usage: python -m services.template_parser <input.docx> [output_dir]")
        sys.exit(1)

    input_file = sys.argv[1]
    out_dir = sys.argv[2] if len(sys.argv) > 2 else "data/blueprints"
    result = parse_docx_to_blueprint(input_file, out_dir)
    print(f"Done → {result}")
